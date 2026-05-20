"""Generate a small docx via python-docx for integration tests.

Run once after `pip install -e .[dev]`; commit the resulting simple.docx.
"""
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]

doc = Document()
doc.add_heading("Hello", level=1)
doc.add_paragraph("This is page one of the test document.")
doc.add_page_break()
doc.add_paragraph("Second page.")
doc.save(Path(__file__).parent / "simple.docx")
